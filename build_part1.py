#!/usr/bin/env python3
"""
Compiles Part I: The Still Water (Prologue + Chapters 1-10) into a printable Word document.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

MANUSCRIPT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "manuscript")
OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Part I - The Still Water.docx")

CHAPTERS = [
    "prologue.md",
    "chapter_01.md",
    "chapter_02.md",
    "chapter_03.md",
    "chapter_04.md",
    "chapter_05.md",
    "chapter_06.md",
    "chapter_07.md",
    "chapter_08.md",
    "chapter_09.md",
    "chapter_10.md",
]

def setup_styles(doc):
    """Configure document styles for a professional manuscript."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.5
    pf.first_line_indent = Inches(0.5)

def add_title_page(doc):
    """Add a title page."""
    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(0)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Inches(0)
    title.paragraph_format.space_after = Pt(24)
    run = title.add_run("MAN AMONGST THE CLOUDS")
    run.font.size = Pt(28)
    run.font.name = 'Georgia'
    run.bold = True

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.first_line_indent = Inches(0)
    sub.paragraph_format.space_after = Pt(36)
    run = sub.add_run("Part I: The Still Water")
    run.font.size = Pt(16)
    run.font.name = 'Georgia'
    run.italic = True

    # Author
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.first_line_indent = Inches(0)
    run = author.add_run("Justin Cronk")
    run.font.size = Pt(16)
    run.font.name = 'Georgia'

    doc.add_page_break()

def add_part_header(doc):
    """Add the Part I divider page."""
    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(0)

    part = doc.add_paragraph()
    part.alignment = WD_ALIGN_PARAGRAPH.CENTER
    part.paragraph_format.first_line_indent = Inches(0)
    part.paragraph_format.space_after = Pt(12)
    run = part.add_run("PART I")
    run.font.size = Pt(14)
    run.font.name = 'Georgia'
    run.bold = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Inches(0)
    title.paragraph_format.space_after = Pt(24)
    run = title.add_run("THE STILL WATER")
    run.font.size = Pt(22)
    run.font.name = 'Georgia'
    run.bold = True

    ep = doc.add_paragraph()
    ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ep.paragraph_format.first_line_indent = Inches(0)
    ep.paragraph_format.space_after = Pt(0)
    run = ep.add_run('"In the time before the taking, the world sang to itself,\nand men were wise enough to listen."')
    run.font.size = Pt(11)
    run.font.name = 'Georgia'
    run.italic = True

    doc.add_page_break()

def add_formatted_text(paragraph, text):
    """Parse markdown italic (*text* or _text_) and add formatted runs."""
    pattern = r'(\*([^*]+)\*|_([^_]+)_)'

    last_end = 0
    for match in re.finditer(pattern, text):
        before = text[last_end:match.start()]
        if before:
            run = paragraph.add_run(before)
            run.font.name = 'Georgia'
            run.font.size = Pt(12)

        italic_text = match.group(2) or match.group(3)
        run = paragraph.add_run(italic_text)
        run.font.name = 'Georgia'
        run.font.size = Pt(12)
        run.italic = True

        last_end = match.end()

    remaining = text[last_end:]
    if remaining:
        run = paragraph.add_run(remaining)
        run.font.name = 'Georgia'
        run.font.size = Pt(12)

def add_chapter(doc, filepath):
    """Parse a chapter markdown file and add it to the document."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    chapter_title = ""
    chapter_subtitle = ""
    body_started = False
    first_para = True

    # Skip "PART ONE" header lines in chapter_01.md
    skip_part_header = True

    for i, line in enumerate(lines):
        stripped = line.strip()

        if not body_started:
            # Skip part header lines (e.g., "# PART ONE: THE STILL WATER")
            if skip_part_header and stripped.startswith('# PART') and 'STILL WATER' in stripped.upper():
                continue
            if skip_part_header and stripped.startswith('_"') and stripped.endswith('"_'):
                continue

            if stripped.startswith('# '):
                chapter_title = stripped.replace('# ', '')
                skip_part_header = False
                continue
            elif stripped.startswith('## '):
                chapter_subtitle = stripped.replace('## ', '')
                continue
            elif stripped == '---':
                if chapter_title:
                    body_started = True
                    # Add spacing before chapter title
                    for _ in range(4):
                        p = doc.add_paragraph()
                        p.paragraph_format.first_line_indent = Inches(0)
                        p.paragraph_format.space_after = Pt(0)

                    # Chapter title
                    ct = doc.add_paragraph()
                    ct.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    ct.paragraph_format.first_line_indent = Inches(0)
                    ct.paragraph_format.space_after = Pt(6)
                    run = ct.add_run(chapter_title.upper())
                    run.font.size = Pt(14)
                    run.font.name = 'Georgia'
                    run.bold = True

                    # Subtitle
                    if chapter_subtitle:
                        st = doc.add_paragraph()
                        st.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        st.paragraph_format.first_line_indent = Inches(0)
                        st.paragraph_format.space_after = Pt(24)
                        run = st.add_run(chapter_subtitle)
                        run.font.size = Pt(12)
                        run.font.name = 'Georgia'
                        run.italic = True

                    first_para = True
                continue
            else:
                continue

        if not body_started:
            continue

        # Scene break
        if stripped == '---':
            scene = doc.add_paragraph()
            scene.alignment = WD_ALIGN_PARAGRAPH.CENTER
            scene.paragraph_format.first_line_indent = Inches(0)
            scene.paragraph_format.space_before = Pt(18)
            scene.paragraph_format.space_after = Pt(18)
            run = scene.add_run("*  *  *")
            run.font.size = Pt(12)
            run.font.name = 'Georgia'
            first_para = True
            continue

        # Skip "End of Part" lines
        if stripped.startswith('*End of Part') or stripped.startswith('_End of Part'):
            continue

        # Empty line
        if stripped == '':
            continue

        # Regular paragraph
        if stripped:
            p = doc.add_paragraph()
            if first_para:
                p.paragraph_format.first_line_indent = Inches(0)
                first_para = False
            else:
                p.paragraph_format.first_line_indent = Inches(0.5)

            add_formatted_text(p, stripped)

    # Page break after chapter
    doc.add_page_break()

def main():
    doc = Document()

    # Page setup — 6x9 trim size for print
    section = doc.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    setup_styles(doc)
    add_title_page(doc)
    add_part_header(doc)

    total_words = 0
    for filename in CHAPTERS:
        filepath = os.path.join(MANUSCRIPT_DIR, filename)
        if not os.path.exists(filepath):
            print(f"WARNING: {filepath} not found, skipping.")
            continue

        print(f"Adding: {filename}")
        add_chapter(doc, filepath)

        with open(filepath, 'r') as f:
            words = len(f.read().split())
            total_words += words

    doc.save(OUTPUT_PATH)
    print(f"\nPart I saved to: {OUTPUT_PATH}")
    print(f"Total: ~{total_words:,} words")

if __name__ == '__main__':
    main()
