#!/usr/bin/env python3
"""
Combines all chapter markdown files into a single formatted Word document.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

MANUSCRIPT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "manuscript")
OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Man Amongst the Clouds - Manuscript.docx")

CHAPTERS = [
    ("prologue.md", None),
    ("chapter_01.md", None),
    ("chapter_02.md", None),
    ("chapter_03.md", None),
    ("chapter_04.md", None),
    ("chapter_05.md", None),
    ("chapter_06.md", None),
    ("chapter_07.md", None),
    ("chapter_08.md", None),
    ("chapter_09.md", None),
    ("chapter_10.md", "END_PART_1"),
    ("chapter_11.md", None),
    ("chapter_12.md", None),
    ("chapter_13.md", None),
    ("chapter_14.md", None),
    ("chapter_15.md", None),
    ("chapter_16.md", None),
    ("chapter_17.md", None),
    ("chapter_18.md", None),
    ("chapter_19.md", None),
    ("chapter_20.md", None),
    ("chapter_21.md", None),
    ("chapter_22.md", "END_PART_2"),
    ("chapter_23.md", None),
    ("chapter_24.md", None),
    ("chapter_25.md", None),
    ("chapter_26.md", None),
    ("chapter_27.md", None),
    ("chapter_28.md", None),
    ("chapter_29.md", None),
    ("chapter_30.md", None),
    ("chapter_31.md", None),
    ("chapter_32.md", None),
    ("chapter_33.md", None),
    ("chapter_34.md", "END_PART_3"),
    ("chapter_35.md", None),
    ("chapter_36.md", None),
    ("chapter_37.md", None),
    ("chapter_38.md", None),
    ("chapter_39.md", None),
    ("chapter_40.md", None),
    ("chapter_41.md", None),
    ("chapter_42.md", "END_PART_4"),
    ("chapter_43.md", None),
    ("chapter_44.md", None),
    ("chapter_45.md", None),
    ("chapter_46.md", None),
    ("chapter_47.md", None),
    ("chapter_48.md", None),
    ("epilogue.md", None),
]

def setup_styles(doc):
    """Configure document styles for a professional manuscript."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.5
    pf.first_line_indent = Inches(0.5)

def add_title_page(doc):
    """Add a title page."""
    # Add blank lines to center vertically
    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(0)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Inches(0)
    title.paragraph_format.space_after = Pt(24)
    run = title.add_run("MAN AMONGST THE CLOUDS")
    run.font.size = Pt(28)
    run.font.name = 'Georgia'
    run.bold = True

    # Subtitle / tagline
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.first_line_indent = Inches(0)
    sub.paragraph_format.space_after = Pt(36)
    run = sub.add_run("A Novel")
    run.font.size = Pt(14)
    run.font.name = 'Georgia'
    run.italic = True

    # Author
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.first_line_indent = Inches(0)
    run = author.add_run("Justin Cronk")
    run.font.size = Pt(16)
    run.font.name = 'Georgia'

    doc.add_page_break()

def add_part_header(doc, part_num, part_title, epigraph=None):
    """Add a part divider page."""
    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(0)

    part = doc.add_paragraph()
    part.alignment = WD_ALIGN_PARAGRAPH.CENTER
    part.paragraph_format.first_line_indent = Inches(0)
    part.paragraph_format.space_after = Pt(12)
    run = part.add_run(f"PART {part_num}")
    run.font.size = Pt(14)
    run.font.name = 'Georgia'
    run.bold = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Inches(0)
    title.paragraph_format.space_after = Pt(24)
    run = title.add_run(part_title)
    run.font.size = Pt(22)
    run.font.name = 'Georgia'
    run.bold = True

    if epigraph:
        ep = doc.add_paragraph()
        ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ep.paragraph_format.first_line_indent = Inches(0)
        ep.paragraph_format.space_after = Pt(0)
        run = ep.add_run(epigraph)
        run.font.size = Pt(11)
        run.font.name = 'Georgia'
        run.italic = True

    doc.add_page_break()

def add_chapter(doc, filepath):
    """Parse a chapter markdown file and add it to the document."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    chapter_title = ""
    chapter_subtitle = ""
    body_started = False
    first_para = True

    for i, line in enumerate(lines):
        stripped = line.strip()

        # Skip empty lines before body starts
        if not body_started:
            if stripped.startswith('# '):
                chapter_title = stripped.replace('# ', '')
                continue
            elif stripped.startswith('## '):
                chapter_subtitle = stripped.replace('## ', '')
                continue
            elif stripped == '---':
                if chapter_title:
                    body_started = True
                    # Add spacing before chapter title
                    for _ in range(4):
                        p = doc.add_paragraph()
                        p.paragraph_format.first_line_indent = Inches(0)
                        p.paragraph_format.space_after = Pt(0)

                    # Chapter title
                    ct = doc.add_paragraph()
                    ct.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    ct.paragraph_format.first_line_indent = Inches(0)
                    ct.paragraph_format.space_after = Pt(6)
                    run = ct.add_run(chapter_title.upper())
                    run.font.size = Pt(14)
                    run.font.name = 'Georgia'
                    run.bold = True

                    # Subtitle
                    if chapter_subtitle:
                        st = doc.add_paragraph()
                        st.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        st.paragraph_format.first_line_indent = Inches(0)
                        st.paragraph_format.space_after = Pt(24)
                        run = st.add_run(chapter_subtitle)
                        run.font.size = Pt(12)
                        run.font.name = 'Georgia'
                        run.italic = True

                    first_para = True
                continue
            else:
                continue

        # Body content
        if not body_started:
            continue

        # Scene break
        if stripped == '---':
            scene = doc.add_paragraph()
            scene.alignment = WD_ALIGN_PARAGRAPH.CENTER
            scene.paragraph_format.first_line_indent = Inches(0)
            scene.paragraph_format.space_before = Pt(18)
            scene.paragraph_format.space_after = Pt(18)
            run = scene.add_run("*  *  *")
            run.font.size = Pt(12)
            run.font.name = 'Georgia'
            first_para = True
            continue

        # Skip "End of Part" lines
        if stripped.startswith('*End of Part') or stripped.startswith('_End of Part'):
            continue

        # Empty line = paragraph separator (skip, handled by paragraph spacing)
        if stripped == '':
            continue

        # Regular paragraph
        if stripped:
            p = doc.add_paragraph()
            if first_para:
                p.paragraph_format.first_line_indent = Inches(0)
                first_para = False
            else:
                p.paragraph_format.first_line_indent = Inches(0.5)

            # Parse inline formatting (italic with * or _)
            add_formatted_text(p, stripped)

    # Page break after chapter
    doc.add_page_break()

def add_formatted_text(paragraph, text):
    """Parse markdown italic (*text* or _text_) and add formatted runs."""
    # Pattern matches *text* or _text_ for italics
    pattern = r'(\*([^*]+)\*|_([^_]+)_)'

    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before the match
        before = text[last_end:match.start()]
        if before:
            run = paragraph.add_run(before)
            run.font.name = 'Georgia'
            run.font.size = Pt(12)

        # Add italic text
        italic_text = match.group(2) or match.group(3)
        run = paragraph.add_run(italic_text)
        run.font.name = 'Georgia'
        run.font.size = Pt(12)
        run.italic = True

        last_end = match.end()

    # Add remaining text
    remaining = text[last_end:]
    if remaining:
        run = paragraph.add_run(remaining)
        run.font.name = 'Georgia'
        run.font.size = Pt(12)

def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    setup_styles(doc)
    add_title_page(doc)

    # Part I header
    add_part_header(doc, "I", "THE STILL WATER",
                    '"In the time before the taking, the world sang to itself,\nand men were wise enough to listen."')

    for filename, flag in CHAPTERS:
        filepath = os.path.join(MANUSCRIPT_DIR, filename)
        if not os.path.exists(filepath):
            print(f"WARNING: {filepath} not found, skipping.")
            continue

        print(f"Adding: {filename}")
        add_chapter(doc, filepath)

        if flag == "END_PART_1":
            # Part II header
            add_part_header(doc, "II", "THE WAKING",
                            '"To learn what you are, first you must lose\neverything you were."')
        elif flag == "END_PART_2":
            # Part III header
            add_part_header(doc, "III", "THE BREAKING",
                            '"The world does not ask if you are ready.\nIt asks if you are willing."')
        elif flag == "END_PART_3":
            # Part IV header
            add_part_header(doc, "IV", "THE CHAMBER",
                            '"He who takes everything possesses nothing.\nHe who gives everything becomes everything."')
        elif flag == "END_PART_4":
            # Part V header
            add_part_header(doc, "V", "THE REMEMBERING",
                            '"The world sings to itself, and men are wise\nenough to listen."')

    doc.save(OUTPUT_PATH)
    print(f"\nManuscript saved to: {OUTPUT_PATH}")

    # Word count estimate
    total_words = 0
    for filename, _ in CHAPTERS:
        filepath = os.path.join(MANUSCRIPT_DIR, filename)
        if os.path.exists(filepath):
            with open(filepath, 'r') as f:
                words = len(f.read().split())
                total_words += words
                print(f"  {filename}: ~{words:,} words")
    print(f"\n  TOTAL: ~{total_words:,} words")

if __name__ == '__main__':
    main()
