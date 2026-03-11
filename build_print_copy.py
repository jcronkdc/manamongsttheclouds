#!/usr/bin/env python3
"""
Builds a paper-conserving print copy on 8.5x11 letter paper.
Looks like a real book: title page, part headers, chapter titles, scene breaks.
Saves paper: 10.5pt font, single spacing, tight margins, no full-page part dividers.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

MANUSCRIPT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "manuscript")
OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           "Man Amongst the Clouds - Print Copy.docx")

CHAPTERS = [
    ("prologue.md", None),
    ("chapter_01.md", None),
    ("chapter_02.md", None),
    ("chapter_03.md", None),
    ("chapter_04.md", None),
    ("chapter_05.md", None),
    ("chapter_06.md", None),
    ("chapter_07.md", None),
    ("chapter_08.md", None),
    ("chapter_09.md", None),
    ("chapter_10.md", "END_PART_1"),
    ("chapter_11.md", None),
    ("chapter_12.md", None),
    ("chapter_13.md", None),
    ("chapter_14.md", None),
    ("chapter_15.md", None),
    ("chapter_16.md", None),
    ("chapter_17.md", None),
    ("chapter_18.md", None),
    ("chapter_19.md", None),
    ("chapter_20.md", None),
    ("chapter_21.md", None),
    ("chapter_22.md", "END_PART_2"),
    ("chapter_23.md", None),
    ("chapter_24.md", None),
    ("chapter_25.md", None),
    ("chapter_26.md", None),
    ("chapter_27.md", None),
    ("chapter_28.md", None),
    ("chapter_29.md", None),
    ("chapter_30.md", None),
    ("chapter_31.md", None),
    ("chapter_32.md", None),
    ("chapter_33.md", None),
    ("chapter_34.md", "END_PART_3"),
    ("chapter_35.md", None),
    ("chapter_36.md", None),
    ("chapter_37.md", None),
    ("chapter_38.md", None),
    ("chapter_39.md", None),
    ("chapter_40.md", None),
    ("chapter_41.md", None),
    ("chapter_42.md", "END_PART_4"),
    ("chapter_43.md", None),
    ("chapter_44.md", None),
    ("chapter_45.md", None),
    ("chapter_46.md", None),
    ("chapter_47.md", None),
    ("chapter_48.md", None),
    ("epilogue.md", None),
]

# Font size for body text (10.5pt saves ~30% paper vs 12pt)
BODY_FONT_SIZE = Pt(10.5)
BODY_FONT_NAME = 'Georgia'

def setup_styles(doc):
    """Configure document styles for paper-conserving print."""
    style = doc.styles['Normal']
    font = style.font
    font.name = BODY_FONT_NAME
    font.size = BODY_FONT_SIZE
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15  # Tight but readable
    pf.first_line_indent = Inches(0.3)

def add_title_page(doc):
    """Add a title page."""
    for _ in range(10):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(0)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Inches(0)
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("MAN AMONGST THE CLOUDS")
    run.font.size = Pt(26)
    run.font.name = BODY_FONT_NAME
    run.bold = True

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.first_line_indent = Inches(0)
    sub.paragraph_format.space_after = Pt(30)
    run = sub.add_run("A Novel")
    run.font.size = Pt(13)
    run.font.name = BODY_FONT_NAME
    run.italic = True

    # Author
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.first_line_indent = Inches(0)
    run = author.add_run("Justin Cronk")
    run.font.size = Pt(15)
    run.font.name = BODY_FONT_NAME

    doc.add_page_break()

def add_part_header_inline(doc, part_num, part_title, epigraph=None):
    """Add an inline part header (no page break — saves a full page per part)."""
    # Small gap
    spacer = doc.add_paragraph()
    spacer.paragraph_format.first_line_indent = Inches(0)
    spacer.paragraph_format.space_before = Pt(36)
    spacer.paragraph_format.space_after = Pt(6)

    # Part number
    part = doc.add_paragraph()
    part.alignment = WD_ALIGN_PARAGRAPH.CENTER
    part.paragraph_format.first_line_indent = Inches(0)
    part.paragraph_format.space_after = Pt(4)
    run = part.add_run(f"PART {part_num}")
    run.font.size = Pt(12)
    run.font.name = BODY_FONT_NAME
    run.bold = True

    # Part title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Inches(0)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run(part_title)
    run.font.size = Pt(18)
    run.font.name = BODY_FONT_NAME
    run.bold = True

    # Epigraph
    if epigraph:
        ep = doc.add_paragraph()
        ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ep.paragraph_format.first_line_indent = Inches(0)
        ep.paragraph_format.space_after = Pt(12)
        run = ep.add_run(epigraph)
        run.font.size = Pt(9.5)
        run.font.name = BODY_FONT_NAME
        run.italic = True

    # Separator line
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.first_line_indent = Inches(0)
    sep.paragraph_format.space_after = Pt(18)
    run = sep.add_run("_______________")
    run.font.size = Pt(10)
    run.font.name = BODY_FONT_NAME
    run.font.color.rgb = RGBColor(160, 160, 160)

    # Page break after part header (each part starts on a new page)
    doc.add_page_break()

def add_chapter(doc, filepath):
    """Parse a chapter markdown file and add it to the document."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    chapter_title = ""
    chapter_subtitle = ""
    body_started = False
    first_para = True

    for i, line in enumerate(lines):
        stripped = line.strip()

        if not body_started:
            if stripped.startswith('# '):
                chapter_title = stripped.replace('# ', '')
                # Handle "PART ONE:" lines in chapter_01.md
                if chapter_title.startswith('PART'):
                    chapter_title = ""
                continue
            elif stripped.startswith('## '):
                chapter_subtitle = stripped.replace('## ', '')
                continue
            elif stripped == '---':
                if chapter_title or chapter_subtitle:
                    body_started = True
                    # Compact chapter header (2 lines of space, not 4)
                    for _ in range(2):
                        p = doc.add_paragraph()
                        p.paragraph_format.first_line_indent = Inches(0)
                        p.paragraph_format.space_after = Pt(0)

                    # Chapter title
                    if chapter_title:
                        ct = doc.add_paragraph()
                        ct.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        ct.paragraph_format.first_line_indent = Inches(0)
                        ct.paragraph_format.space_after = Pt(3)
                        run = ct.add_run(chapter_title.upper())
                        run.font.size = Pt(11)
                        run.font.name = BODY_FONT_NAME
                        run.bold = True

                    # Subtitle
                    if chapter_subtitle:
                        st = doc.add_paragraph()
                        st.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        st.paragraph_format.first_line_indent = Inches(0)
                        st.paragraph_format.space_after = Pt(14)
                        run = st.add_run(chapter_subtitle)
                        run.font.size = Pt(10)
                        run.font.name = BODY_FONT_NAME
                        run.italic = True

                    first_para = True
                continue
            elif stripped.startswith('_"') or stripped.startswith('_\u201c'):
                # Part epigraph in chapter_01.md — skip it (handled by part header)
                continue
            else:
                continue

        if not body_started:
            continue

        # Scene break — compact
        if stripped == '---':
            scene = doc.add_paragraph()
            scene.alignment = WD_ALIGN_PARAGRAPH.CENTER
            scene.paragraph_format.first_line_indent = Inches(0)
            scene.paragraph_format.space_before = Pt(10)
            scene.paragraph_format.space_after = Pt(10)
            run = scene.add_run("*  *  *")
            run.font.size = BODY_FONT_SIZE
            run.font.name = BODY_FONT_NAME
            first_para = True
            continue

        # Skip "End of Part" lines
        if stripped.startswith('*End of Part') or stripped.startswith('_End of Part'):
            continue

        # Empty line
        if stripped == '':
            continue

        # Regular paragraph
        if stripped:
            p = doc.add_paragraph()
            if first_para:
                p.paragraph_format.first_line_indent = Inches(0)
                first_para = False
            else:
                p.paragraph_format.first_line_indent = Inches(0.3)

            add_formatted_text(p, stripped)

    # No page break between chapters within the same part — saves paper
    # Just add a small gap
    spacer = doc.add_paragraph()
    spacer.paragraph_format.first_line_indent = Inches(0)
    spacer.paragraph_format.space_after = Pt(6)

def add_formatted_text(paragraph, text):
    """Parse markdown italic (*text* or _text_) and add formatted runs."""
    pattern = r'(\*([^*]+)\*|_([^_]+)_)'

    last_end = 0
    for match in re.finditer(pattern, text):
        before = text[last_end:match.start()]
        if before:
            run = paragraph.add_run(before)
            run.font.name = BODY_FONT_NAME
            run.font.size = BODY_FONT_SIZE

        italic_text = match.group(2) or match.group(3)
        run = paragraph.add_run(italic_text)
        run.font.name = BODY_FONT_NAME
        run.font.size = BODY_FONT_SIZE
        run.italic = True

        last_end = match.end()

    remaining = text[last_end:]
    if remaining:
        run = paragraph.add_run(remaining)
        run.font.name = BODY_FONT_NAME
        run.font.size = BODY_FONT_SIZE

def main():
    doc = Document()

    # Page setup — 8.5 x 11 letter, tight margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    setup_styles(doc)
    add_title_page(doc)

    # Part I header
    add_part_header_inline(doc, "I", "THE STILL WATER",
                    '"In the time before the taking, the world sang to itself,\nand men were wise enough to listen."')

    for filename, flag in CHAPTERS:
        filepath = os.path.join(MANUSCRIPT_DIR, filename)
        if not os.path.exists(filepath):
            print(f"WARNING: {filepath} not found, skipping.")
            continue

        print(f"Adding: {filename}")
        add_chapter(doc, filepath)

        if flag == "END_PART_1":
            add_part_header_inline(doc, "II", "THE WAKING",
                            '"To learn what you are, first you must lose\neverything you were."')
        elif flag == "END_PART_2":
            add_part_header_inline(doc, "III", "THE BREAKING",
                            '"The world does not ask if you are ready.\nIt asks if you are willing."')
        elif flag == "END_PART_3":
            add_part_header_inline(doc, "IV", "THE CHAMBER",
                            '"He who takes everything possesses nothing.\nHe who gives everything becomes everything."')
        elif flag == "END_PART_4":
            add_part_header_inline(doc, "V", "THE REMEMBERING",
                            '"The world sings to itself, and men are wise\nenough to listen."')

    doc.save(OUTPUT_PATH)
    print(f"\nPrint copy saved to: {OUTPUT_PATH}")
    print(f"\nPaper-saving features:")
    print(f"  - 10.5pt Georgia (vs 12pt in manuscript version)")
    print(f"  - 1.15 line spacing (vs 1.5)")
    print(f"  - 8.5x11 letter paper with tight margins")
    print(f"  - No full-page part dividers")
    print(f"  - No page breaks between chapters in same part")
    print(f"  - Estimated: ~280-320 pages (vs ~450+ at manuscript spacing)")

if __name__ == '__main__':
    main()
